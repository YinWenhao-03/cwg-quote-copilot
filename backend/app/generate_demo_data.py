from __future__ import annotations

import json
from datetime import UTC, date, datetime, timedelta
from email.message import EmailMessage
from pathlib import Path
from random import Random

from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

from .config import ROOT

PRODUCTS = [
    {
        "sku": f"S4-{1000 + index}",
        "name": f"车载智能座舱显示模组 {1000 + index}",
        "weight_kg": round(0.8 + index * 0.07, 2),
        "packaging_cost": round(12 + index * 0.8, 2),
        "list_price": round(980 + index * 42, 2),
    }
    for index in range(20)
]

CUSTOMERS = [
    {"id": "CUST-BMW-CN", "name": "华东汽车科技", "destination": "上海", "incoterm": "DDP"},
    {"id": "CUST-GEELY", "name": "远航智能汽车", "destination": "宁波", "incoterm": "DAP"},
    {"id": "CUST-EU-DE", "name": "Europa Mobility GmbH", "destination": "汉堡", "incoterm": "DDP"},
]


def _write_pdf(path: Path, title: str, lines: list[str]) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    c = canvas.Canvas(str(path), pagesize=A4)
    _, height = A4
    c.setFont("STSong-Light", 16)
    c.drawString(56, height - 60, title)
    c.setFont("STSong-Light", 10)
    y = height - 90
    for line in lines:
        if y < 60:
            c.showPage()
            c.setFont("STSong-Light", 10)
            y = height - 60
        c.drawString(56, y, line[:85])
        y -= 18
    c.save()


def _write_docx(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    doc = WordDocument()
    doc.add_heading(title, 0)
    for heading, content in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content)
    doc.save(path)


def generate_demo_files(output_dir: Path | None = None) -> list[dict[str, object]]:
    output_dir = output_dir or ROOT / "data" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    rng = Random(20260811)
    manifest: list[dict[str, object]] = []

    for index, product in enumerate(PRODUCTS):
        sku = str(product["sku"])
        pdf_path = output_dir / f"product_{sku}.pdf"
        _write_pdf(
            pdf_path,
            f"{sku} 产品技术资料",
            [
                f"产品名称：{product['name']}",
                f"净重：{product['weight_kg']} kg",
                "显示接口：车规级高速视频接口，支持温度补偿。",
                "包装要求：每10件使用防静电内袋并装入加强纸箱。",
                "质量标准：出货前执行外观、亮度、接口和高低温抽检。",
                "本文档仅适用于当前已批准版本，历史版本不得用于报价。",
            ],
        )
        manifest.append(
            {
                "path": str(pdf_path),
                "title": f"{sku} 产品技术资料",
                "document_type": "product_manual",
                "classification": "sales",
                "sku": sku,
                "status": "approved",
                "valid_from": "2026-01-01",
                "valid_to": "2027-12-31",
            }
        )

        docx_path = output_dir / f"quality_guide_{sku}.docx"
        _write_docx(
            docx_path,
            f"{sku} 质量与包装指南",
            [
                ("包装", "国内运输采用十件一箱；出口运输增加防潮袋、边角保护和跌落标签。"),
                ("检验", "亮度、色差和高低温循环均需记录批次号并保留十二个月。"),
                ("异常处理", "批量异常须在二十四小时内提交8D问题单，不得由销售自行承诺赔付。"),
            ],
        )
        manifest.append(
            {
                "path": str(docx_path),
                "title": f"{sku} 质量与包装指南",
                "document_type": "quality_guide",
                "classification": "sales",
                "sku": sku,
                "status": "approved",
                "valid_from": "2026-02-01",
                "valid_to": "2027-12-31",
            }
        )

    for index in range(25):
        product = PRODUCTS[index % len(PRODUCTS)]
        customer = CUSTOMERS[index % len(CUSTOMERS)]
        quote_date = date(2025, 1, 10) + timedelta(days=index * 11)
        price = float(product["list_price"]) * rng.uniform(0.92, 1.08)
        path = output_dir / f"historical_quote_{index + 1:03d}.pdf"
        _write_pdf(
            path,
            f"历史报价 QT-2025-{index + 1:03d}",
            [
                f"客户：{customer['name']}",
                f"产品：{product['sku']}",
                f"数量：{200 + index * 20}",
                f"单价：{price:.2f} CNY",
                f"贸易条款：{customer['incoterm']}",
                f"报价日期：{quote_date.isoformat()}",
                "注意：本文件为历史报价，仅用于趋势比较，不代表当前有效价格。",
            ],
        )
        manifest.append(
            {
                "path": str(path),
                "title": f"历史报价 QT-2025-{index + 1:03d}",
                "document_type": "historical_quote",
                "classification": "sales",
                "sku": product["sku"],
                "customer_id": customer["id"],
                "status": "approved",
                "valid_from": quote_date.isoformat(),
                "valid_to": quote_date.isoformat(),
            }
        )

    for index in range(15):
        customer = CUSTOMERS[index % len(CUSTOMERS)]
        path = output_dir / f"contract_{index + 1:03d}.docx"
        _write_docx(
            path,
            f"{customer['name']} 采购框架合同 {index + 1:03d}",
            [
                ("1. 适用范围", "合同适用于双方确认的车载显示模组及配套服务。"),
                (
                    "2. 交付",
                    f"默认目的地为{customer['destination']}，贸易条款为{customer['incoterm']}。",
                ),
                ("3. 报价有效期", "正式报价自批准日起三十日内有效，变更需重新审批。"),
                ("4. 保密", "供应商成本、内部底价和利润规则不得向客户或普通销售披露。"),
            ],
        )
        manifest.append(
            {
                "path": str(path),
                "title": f"{customer['name']} 采购框架合同",
                "document_type": "contract",
                "classification": "sales",
                "customer_id": customer["id"],
                "status": "approved",
                "valid_from": "2026-01-01",
                "valid_to": "2027-12-31",
            }
        )

    management_path = output_dir / "management_price_policy.pdf"
    _write_pdf(
        management_path,
        "管理层内部价格政策",
        [
            "密级：管理层机密。",
            "内部底价、硬毛利底线和供应商成本不得进入销售人员的检索结果或模型上下文。",
            "低于标准最低价的例外报价必须由经理填写理由并保留审计记录。",
            "任何低于硬底价的报价均应由系统直接阻断。",
        ],
    )
    manifest.append(
        {
            "path": str(management_path),
            "title": "管理层内部价格政策",
            "document_type": "pricing_policy",
            "classification": "management",
            "status": "approved",
            "valid_from": "2026-01-01",
            "valid_to": "2027-12-31",
        }
    )

    for index in range(15):
        product = PRODUCTS[index % len(PRODUCTS)]
        path = output_dir / f"quality_report_{index + 1:03d}.pdf"
        _write_pdf(
            path,
            f"{product['sku']} 质量报告 QR-{index + 1:03d}",
            [
                f"批次：BATCH-2026-{index + 1:03d}",
                "外观抽检：通过",
                "亮度均匀性：通过",
                "高低温循环：通过",
                "接口稳定性：通过",
                "结论：允许出货。",
            ],
        )
        manifest.append(
            {
                "path": str(path),
                "title": f"{product['sku']} 质量报告",
                "document_type": "quality_report",
                "classification": "sales",
                "sku": product["sku"],
                "status": "approved",
                "valid_from": "2026-06-01",
                "valid_to": "2027-05-31",
            }
        )

    for index in range(10):
        product = PRODUCTS[index]
        path = output_dir / f"supplier_cost_{index + 1:03d}.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "Approved Costs"
        ws.append(["SKU", "Supplier", "Unit Cost", "Currency", "Valid From", "Valid To", "Status"])
        ws.append(
            [
                product["sku"],
                f"Supplier-{index % 3 + 1}",
                650 + index * 23,
                "CNY",
                "2026-07-01",
                "2026-12-31",
                "approved",
            ]
        )
        ws.append(
            [
                product["sku"],
                f"Supplier-{index % 3 + 1}",
                590 + index * 20,
                "CNY",
                "2025-01-01",
                "2025-12-31",
                "expired",
            ]
        )
        wb.save(path)
        manifest.append(
            {
                "path": str(path),
                "title": f"{product['sku']} 供应商成本",
                "document_type": "supplier_cost",
                "classification": "procurement",
                "sku": product["sku"],
                "status": "approved",
                "valid_from": "2026-07-01",
                "valid_to": "2026-12-31",
            }
        )

    for index in range(15):
        product = PRODUCTS[index % len(PRODUCTS)]
        customer = CUSTOMERS[index % len(CUSTOMERS)]
        message = EmailMessage()
        message["From"] = f"buyer{index + 1}@example.test"
        message["To"] = "sales@cwg.local"
        message["Subject"] = f"询价：{product['sku']} {300 + index * 20}件"
        message["Date"] = datetime(2026, 8, 1 + index % 10, 9, 0, tzinfo=UTC).strftime(
            "%a, %d %b %Y %H:%M:%S %z"
        )
        currency_text = "币种CNY，" if index % 4 else ""
        message.set_content(
            f"你好，我们是{customer['name']}。请对{product['sku']}报价{300 + index * 20}件，"
            f"纸箱包装，发往{customer['destination']}，贸易条款{customer['incoterm']}，"
            f"{currency_text}希望九月底前交付。谢谢。"
        )
        path = output_dir / f"inquiry_{index + 1:03d}.eml"
        path.write_bytes(message.as_bytes())
        manifest.append(
            {
                "path": str(path),
                "title": str(message["Subject"]),
                "document_type": "email",
                "classification": "sales",
                "sku": product["sku"],
                "customer_id": customer["id"],
                "status": "approved",
                "valid_from": "2026-08-01",
                "valid_to": "2027-08-01",
                "inbox": True,
            }
        )

    (output_dir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return manifest


if __name__ == "__main__":
    entries = generate_demo_files()
    print(f"Generated {len(entries)} demo files.")
