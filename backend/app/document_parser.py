from __future__ import annotations

import hashlib
import re
import uuid
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from pathlib import Path

from docx import Document as WordDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .config import get_settings


@dataclass(slots=True)
class ParsedBlock:
    content: str
    page: int | None = None
    section: str | None = None


class DocumentParser:
    def parse(self, path: Path) -> list[ParsedBlock]:
        suffix = path.suffix.lower()
        if get_settings().document_parser == "docling" and suffix in {".pdf", ".docx", ".xlsx"}:
            return self._parse_docling(path)
        if suffix == ".pdf":
            return self._parse_pdf(path)
        if suffix == ".docx":
            return self._parse_docx(path)
        if suffix == ".xlsx":
            return self._parse_xlsx(path)
        if suffix in {".eml", ".msg"}:
            return self._parse_email(path)
        return [ParsedBlock(path.read_text(encoding="utf-8", errors="ignore"))]

    def _parse_docling(self, path: Path) -> list[ParsedBlock]:
        try:
            from docling.document_converter import DocumentConverter
        except ImportError as exc:
            raise RuntimeError("DOCUMENT_PARSER=docling 需要安装 backend 的 full 依赖") from exc
        result = DocumentConverter().convert(path)
        markdown = result.document.export_to_markdown()
        return self._paragraph_blocks(markdown)

    def _parse_pdf(self, path: Path) -> list[ParsedBlock]:
        blocks: list[ParsedBlock] = []
        reader = PdfReader(str(path))
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            blocks.extend(self._paragraph_blocks(text, page=page_number))
        return blocks

    def _parse_docx(self, path: Path) -> list[ParsedBlock]:
        doc = WordDocument(path)
        blocks: list[ParsedBlock] = []
        current_section: str | None = None
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style.name.startswith("Heading") or paragraph.style.name == "Title":
                current_section = text
            else:
                blocks.append(ParsedBlock(text, section=current_section))
        for table_index, table in enumerate(doc.tables, start=1):
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            blocks.append(ParsedBlock("\n".join(rows), section=f"表格 {table_index}"))
        return blocks

    def _parse_xlsx(self, path: Path) -> list[ParsedBlock]:
        workbook = load_workbook(path, data_only=False, read_only=True)
        blocks: list[ParsedBlock] = []
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(value or "") for value in rows[0]]
            for row_number, row in enumerate(rows[1:], start=2):
                values = [str(value or "") for value in row]
                content = "；".join(
                    f"{header}：{value}"
                    for header, value in zip(headers, values, strict=False)
                    if value
                )
                if content:
                    blocks.append(ParsedBlock(content, section=f"{sheet.title}!{row_number}"))
        return blocks

    def _parse_email(self, path: Path) -> list[ParsedBlock]:
        message = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
        body = message.get_body(preferencelist=("plain",))
        text = body.get_content() if body else ""
        header = f"主题：{message.get('Subject', '')}\n发件人：{message.get('From', '')}"
        return [ParsedBlock(f"{header}\n{text}", section="邮件正文")]

    @staticmethod
    def _paragraph_blocks(text: str, page: int | None = None) -> list[ParsedBlock]:
        paragraphs = [
            part.strip() for part in re.split(r"\n{2,}|(?<=[。；])\s*", text) if part.strip()
        ]
        return [ParsedBlock(paragraph, page=page) for paragraph in paragraphs]


def chunk_blocks(
    blocks: list[ParsedBlock], *, max_chars: int = 700, overlap_chars: int = 80
) -> list[ParsedBlock]:
    chunks: list[ParsedBlock] = []
    buffer = ""
    page: int | None = None
    section: str | None = None
    for block in blocks:
        if buffer and len(buffer) + len(block.content) + 1 > max_chars:
            chunks.append(ParsedBlock(buffer.strip(), page=page, section=section))
            buffer = buffer[-overlap_chars:] if overlap_chars else ""
        if not buffer:
            page = block.page
            section = block.section
        buffer = f"{buffer}\n{block.content}".strip()
    if buffer:
        chunks.append(ParsedBlock(buffer.strip(), page=page, section=section))
    return chunks


def stable_chunk_id(version_id: str, ordinal: int, content: str) -> str:
    raw = f"{version_id}:{ordinal}:{content}".encode()
    return str(uuid.UUID(bytes=hashlib.sha256(raw).digest()[:16]))
